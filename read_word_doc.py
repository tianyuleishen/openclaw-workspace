#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取 Word 文档并保存为文本
"""

from docx import Document
import sys

def read_word_doc(doc_path, max_chars=50000):
    """读取 Word 文档，返回纯文本"""
    print(f"📖 正在读取文档: {doc_path}")
    print(f"文件大小: {__import__('os').path.getsize(doc_path) / 1024 / 1024:.2f} MB")
    
    doc = Document(doc_path)
    
    all_text = []
    
    # 读取段落
    print("📝 读取段落...")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # 读取表格
    print("📊 读取表格...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_text.append(text)
    
    full_text = '\n\n'.join(all_text)
    
    print(f"✅ 读取完成！总共 {len(full_text)} 字符")
    
    return full_text

if __name__ == '__main__':
    doc_path = '/home/开发.doc'
    
    try:
        text = read_word_doc(doc_path)
        
        # 保存为文本文件
        output_path = '/home/admin/.openclaw/workspace/docs/开发_doc_content.txt'
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        print(f"\n💾 已保存到: {output_path}")
        print(f"📏 文件大小: {len(text)} 字符")
        
        # 显示前 2000 个字符预览
        print("\n" + "="*60)
        print("📖 文档内容预览（前2000字符）：")
        print("="*60)
        print(text[:2000])
        print("...\n" + "="*60)
        
    except Exception as e:
        print(f"❌ 读取失败: {e}")
        sys.exit(1)
